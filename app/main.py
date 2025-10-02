from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse
from pydantic import BaseModel
from typing import List, Optional
import os
import asyncio
from concurrent.futures import ThreadPoolExecutor
import traceback

class ChatRequest(BaseModel):
    question: str
    context: Optional[str] = None

class SourceDocument(BaseModel):
    title: str
    doc_id: str

class ChatResponse(BaseModel):
    question: str
    answer: str
    context: Optional[List[str]] = None
    sources: Optional[List[SourceDocument]] = None

class Document(BaseModel):
    title: str
    content: str

app = FastAPI(title="GraphRAG Chatbot", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploaded_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Thread pool for parallel processing
executor = ThreadPoolExecutor(max_workers=4)

rag_system = None
postgres = None
file_storage = {}

print("Initializing GraphRAG system...")
try:
    from rag import GraphRAGSystem
    rag_system = GraphRAGSystem()
    print("GraphRAG system initialized")
except Exception as e:
    print(f"GraphRAG initialization failed: {e}")

try:
    from database import PostgresConnection
    print("Connecting to PostgreSQL...")
    postgres = PostgresConnection()
    postgres.connect()
    print("PostgreSQL connected!")
except Exception as e:
    print(f"PostgreSQL connection failed: {e}")
    postgres = None

@app.get("/")
async def root():
    return {"message": "GraphRAG Chatbot API is running!"}

@app.get("/status")
async def get_status():
    neo4j_status = "disconnected"
    if rag_system and hasattr(rag_system, 'neo4j'):
        try:
            if hasattr(rag_system.neo4j, 'driver') and rag_system.neo4j.driver:
                with rag_system.neo4j.driver.session() as session:
                    session.run("RETURN 1")
                neo4j_status = "connected"
        except Exception:
            neo4j_status = "disconnected"
    
    postgres_status = "disconnected"
    if postgres and postgres.conn:
        try:
            with postgres.conn.cursor() as cursor:
                cursor.execute("SELECT 1")
            postgres_status = "connected"
        except Exception:
            postgres_status = "disconnected"
    
    ollama_status = "disconnected"
    model_name = "unknown"
    try:
        import ollama
        ollama.list()
        ollama_status = "connected"
        if rag_system and hasattr(rag_system, 'model_name'):
            model_name = rag_system.model_name
    except Exception:
        pass
    
    return {
        "message": "GraphRAG Chatbot API is running!",
        "neo4j": neo4j_status,
        "postgresql": postgres_status,
        "ollama": ollama_status,
        "model": model_name
    }

@app.get("/documents")
async def list_documents():
    if not rag_system:
        return {"documents": [], "count": 0, "message": "GraphRAG system not available"}
    
    try:
        if hasattr(rag_system, 'neo4j') and hasattr(rag_system.neo4j, 'driver') and rag_system.neo4j.driver:
            with rag_system.neo4j.driver.session() as session:
                result = session.run("MATCH (d:Document) RETURN d.title as title, d.id as id ORDER BY d.title")
                documents = [{"id": record["id"], "title": record["title"]} for record in result]
            return {"documents": documents, "count": len(documents)}
        else:
            return {"documents": [], "count": 0, "message": "Neo4j not connected"}
    except Exception as e:
        print(f"Error listing documents: {e}")
        return {"documents": [], "count": 0, "error": str(e)}

@app.get("/document/{doc_id}")
async def get_document(doc_id: str):
    """Get document content for viewing in browser"""
    if doc_id in file_storage:
        file_path = file_storage[doc_id]
        if os.path.exists(file_path):
            file_ext = os.path.splitext(file_path.lower())[1]
            
            if file_ext in ['.txt', '.md']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    return PlainTextResponse(
                        content=content,
                        headers={
                            "Content-Type": "text/plain; charset=utf-8",
                            "Content-Disposition": f"inline; filename={os.path.basename(file_path)}"
                        }
                    )
                except Exception as e:
                    raise HTTPException(status_code=500, detail=f"Error reading file: {str(e)}")
            
            elif file_ext == '.pdf':
                return FileResponse(
                    file_path,
                    media_type='application/pdf',
                    headers={
                        "Content-Disposition": f"inline; filename={os.path.basename(file_path)}"
                    }
                )
            
            elif file_ext == '.docx':
                raise HTTPException(
                    status_code=415, 
                    detail="DOCX files cannot be viewed directly in browser. The document content was used for the answer."
                )
            
            else:
                return FileResponse(
                    file_path,
                    media_type='application/octet-stream',
                    headers={
                        "Content-Disposition": f"inline; filename={os.path.basename(file_path)}"
                    }
                )
    
    raise HTTPException(status_code=404, detail="Document file not found")

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Optimized text extraction with better error handling"""
    file_ext = os.path.splitext(filename.lower())[1]
    
    if file_ext in ['.txt', '.md']:
        # Try UTF-8 first (most common)
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return content.decode(encoding, errors='ignore')
                except:
                    continue
        raise ValueError("Could not decode text file")
    
    elif file_ext == '.pdf':
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Parallel page extraction for large PDFs
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            
            return "\n".join(text_parts).strip()
        except ImportError:
            raise ValueError("PyPDF2 not installed. Install with: pip install PyPDF2")
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    elif file_ext == '.docx':
        try:
            import docx
            from io import BytesIO
            
            doc_file = BytesIO(content)
            doc = docx.Document(doc_file)
            
            # Extract all paragraphs efficiently
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            return text.strip()
        except ImportError:
            raise ValueError("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

async def process_single_file(file: UploadFile):
    """Process a single file asynchronously"""
    try:
        if not file.filename:
            return {"success": False, "name": "unnamed_file", "reason": "No filename"}
        
        file_ext = os.path.splitext(file.filename.lower())[1]
        allowed_exts = {'.txt', '.md', '.pdf', '.docx'}
        
        if file_ext not in allowed_exts:
            return {
                "success": False,
                "name": file.filename,
                "reason": f"Unsupported format: {file_ext}"
            }
        
        # Read file content
        content = await file.read()
        if not content:
            return {"success": False, "name": file.filename, "reason": "Empty file"}
        
        # Save file
        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # Extract text (run in thread pool to avoid blocking)
        loop = asyncio.get_event_loop()
        try:
            text_content = await loop.run_in_executor(
                executor, 
                extract_text_from_file, 
                content, 
                file.filename
            )
            
            if not text_content or len(text_content.strip()) < 10:
                return {"success": False, "name": file.filename, "reason": "No readable text"}
            
        except ValueError as e:
            return {"success": False, "name": file.filename, "reason": str(e)}
        except Exception as e:
            return {"success": False, "name": file.filename, "reason": f"Parse error: {str(e)}"}
        
        # Store in RAG system (run in thread pool)
        try:
            await loop.run_in_executor(
                executor,
                rag_system.store_document,
                file.filename,
                text_content
            )
            
            # Get document ID
            doc_id = None
            if hasattr(rag_system, 'neo4j') and rag_system.neo4j.driver:
                with rag_system.neo4j.driver.session() as session:
                    result = session.run(
                        "MATCH (d:Document {title: $title}) RETURN d.id as id",
                        title=file.filename
                    )
                    record = result.single()
                    if record:
                        doc_id = record["id"]
                        file_storage[doc_id] = file_path
            
            # PostgreSQL backup (non-blocking)
            if postgres and postgres.conn:
                try:
                    with postgres.conn.cursor() as cursor:
                        cursor.execute(
                            "INSERT INTO documents (title, content) VALUES (%s, %s)",
                            (file.filename, text_content[:5000])
                        )
                        postgres.conn.commit()
                except Exception as e:
                    print(f"PostgreSQL backup warning: {e}")
            
            return {"success": True, "name": file.filename}
            
        except Exception as e:
            return {"success": False, "name": file.filename, "reason": f"Storage failed: {str(e)}"}
            
    except Exception as e:
        return {"success": False, "name": file.filename if file.filename else "unknown", "reason": f"Processing error: {str(e)}"}

@app.post("/upload-files")
async def upload_files(files: List[UploadFile] = File(...)):
    """Optimized parallel file upload"""
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    if not rag_system:
        raise HTTPException(status_code=503, detail="GraphRAG system not available")
    
    # Process all files in parallel
    tasks = [process_single_file(file) for file in files]
    results = await asyncio.gather(*tasks)
    
    # Separate successful and failed uploads
    successful = [r["name"] for r in results if r["success"]]
    failed = [{"name": r["name"], "reason": r["reason"]} for r in results if not r["success"]]
    
    if not successful and failed:
        error_summary = "; ".join([f"{f['name']}: {f['reason']}" for f in failed[:3]])
        raise HTTPException(status_code=400, detail=f"All files failed. Errors: {error_summary}")
    
    message = f"Successfully processed {len(successful)}/{len(files)} file(s)"
    if failed:
        message += f". {len(failed)} file(s) failed"
    
    return {
        "message": message,
        "processed_files": successful,
        "failed_files": failed,
        "success_count": len(successful),
        "failure_count": len(failed)
    }

@app.post("/chat", response_model=ChatResponse)
async def chat_endpoint(request: ChatRequest):
    if not rag_system:
        raise HTTPException(status_code=503, detail="GraphRAG system not available")
    
    try:
        print(f"Chat request: {request.question}")
        
        # Run chat in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            executor,
            rag_system.chat,
            request.question
        )
        
        print(f"Chat result keys: {result.keys()}")
        print(f"Sources in result: {result.get('sources', [])}")
        
        if postgres and postgres.conn:
            try:
                with postgres.conn.cursor() as cursor:
                    cursor.execute(
                        "INSERT INTO chat_history (question, answer) VALUES (%s, %s)",
                        (result["question"], result["answer"])
                    )
                    postgres.conn.commit()
            except Exception as e:
                print(f"Chat history save failed: {e}")
        
        sources = []
        if result.get("sources"):
            sources = [SourceDocument(title=s["title"], doc_id=s["doc_id"]) for s in result["sources"]]
            print(f"Returning {len(sources)} sources to frontend")
        
        response = ChatResponse(
            question=result["question"],
            answer=result["answer"],
            context=result.get("context", []),
            sources=sources
        )
        
        return response
        
    except Exception as e:
        print(f"Chat error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")

@app.delete("/documents")
async def clear_documents():
    if not rag_system:
        raise HTTPException(status_code=503, detail="GraphRAG system not available")
    
    try:
        if hasattr(rag_system, 'neo4j') and rag_system.neo4j.driver:
            with rag_system.neo4j.driver.session() as session:
                count_result = session.run("MATCH (d:Document) RETURN count(d) as count")
                doc_count = count_result.single()["count"]
                
                session.run("MATCH (n) DETACH DELETE n")
                
                if postgres and postgres.conn:
                    with postgres.conn.cursor() as cursor:
                        cursor.execute("DELETE FROM documents")
                        postgres.conn.commit()
                
                for file_path in file_storage.values():
                    if os.path.exists(file_path):
                        os.remove(file_path)
                file_storage.clear()
                
                return {"message": f"Cleared {doc_count} documents"}
        else:
            return {"message": "No Neo4j connection to clear"}
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Clear failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")